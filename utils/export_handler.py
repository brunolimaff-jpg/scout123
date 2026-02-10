import io
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.units import cm
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json

class ExportHandler:
    
    @staticmethod
    def generate_pdf(dossie_data):
        """Gera PDF profissional 'violento' com todas as 5 camadas."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=1.5*cm, leftMargin=1.5*cm,
                                topMargin=1.5*cm, bottomMargin=1.5*cm)
        
        elements = []
        styles = getSampleStyleSheet()
        
        # ===== ESTILOS =====
        title_style = ParagraphStyle(
            'TitleCustom', parent=styles['Heading1'], fontSize=22, spaceAfter=12, 
            textColor=colors.HexColor("#0f172a"), alignment=1, fontName='Helvetica-Bold'
        )
        heading2_style = ParagraphStyle(
            'Heading2Custom', parent=styles['Heading2'], fontSize=14, spaceAfter=8,
            textColor=colors.HexColor("#1e40af"), fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'BodyCustom', parent=styles['Normal'], fontSize=9, spaceAfter=6, leading=12
        )
        alert_style = ParagraphStyle(
            'AlertCustom', parent=styles['Normal'], fontSize=9, textColor=colors.red,
            spaceAfter=4, fontName='Helvetica-Bold'
        )
        
        # ===== CABEÇALHO =====
        elements.append(Paragraph("🔴 RADAR FOX-3 | INTELLIGENCE REPORT", title_style))
        elements.append(Spacer(1, 0.3*cm))
        empresa_alvo = dossie_data.get('empresa_alvo', 'ALVO DESCONHECIDO')
        sas_score = dossie_data.get('sas_score', 'N/A')
        sas_tier = dossie_data.get('sas_tier', 'N/A')
        elements.append(Paragraph(f"<b>ALVO:</b> {empresa_alvo.upper()} | <b>SCORE SAS:</b> {sas_score} ({sas_tier}) | <b>DATA:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", body_style))
        elements.append(Spacer(1, 0.5*cm))
        
        # ===== SCORECARD RESUMIDO =====
        elementos_resumo = [
            ['INDICADOR', 'DADOS RECUPERADOS'],
            ['Área Total', f"{dossie_data.get('dados_operacionais', {}).get('area_total', 'N/D')} ha"],
            ['Faturamento Estimado', dossie_data.get('dados_financeiros', {}).get('faturamento_estimado', 'N/D')],
            ['EBITDA Ajustado', dossie_data.get('dados_financeiros', {}).get('ebitda_ajustado', 'N/D')],
            ['Dívida Total', dossie_data.get('dados_financeiros', {}).get('divida_total', 'N/D')],
            ['ERP Principal', dossie_data.get('tech_stack', {}).get('erp_principal', 'N/D')],
            ['Funcionários Estimados', dossie_data.get('dados_operacionais', {}).get('funcionarios_estimados', 'N/D')]
        ]
        
        table_resumo = Table(elementos_resumo, colWidths=[4*cm, 8*cm])
        table_resumo.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        elements.append(table_resumo)
        elements.append(Spacer(1, 0.5*cm))
        
        # ===== ALERTAS CRÍTICOS (se houver) =====
        validacoes = dossie_data.get('validacoes_criticas', {})
        if validacoes.get('alertas'):
            elements.append(Paragraph("⚠️ ALERTAS CRÍTICOS", heading2_style))
            for alerta in validacoes.get('alertas', []):
                elementos_alerta = [
                    Paragraph(f"<b>[{alerta.get('tipo')}]</b> - Severidade: {alerta.get('severidade')}", alert_style),
                    Paragraph(f"Mensagem: {alerta.get('mensagem')}", body_style),
                    Paragraph(f"Ação: {alerta.get('acao')}", body_style)
                ]
                for e in elementos_alerta:
                    elements.append(e)
                elements.append(Spacer(1, 0.2*cm))
            elements.append(Spacer(1, 0.4*cm))
        
        # ===== SEÇÃO 1: INFRAESTRUTURA (Hard Assets) =====
        elements.append(PageBreak())
        elements.append(Paragraph("1️⃣ CAMADA DE INFRAESTRUTURA (HARD ASSETS)", heading2_style))
        elements.append(Spacer(1, 0.3*cm))
        
        # CAR/SIGEF
        car_data = dossie_data.get('sigef_car', {})
        elements.append(Paragraph(f"<b>SIGEF/CAR:</b> {car_data.get('area_total_hectares', 'N/D')} ha mapeados | Regularização: {car_data.get('regularizacao_percentual', 'N/D')}%", body_style))
        
        car_registros = car_data.get('car_records', [])
        if car_registros:
            car_table_data = [['Município', 'UF', 'Hectares', 'Status']]
            for reg in car_registros[:5]:  # Primeiras 5
                car_table_data.append([
                    reg.get('municipio', '-'),
                    reg.get('uf', '-'),
                    f"{reg.get('hectares', 0):,}",
                    reg.get('status', '-')
                ])
            car_table = Table(car_table_data, colWidths=[3.5*cm, 1.5*cm, 2*cm, 2*cm])
            car_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            elements.append(car_table)
        
        elements.append(Spacer(1, 0.2*cm))
        
        # Frota
        frota_data = dossie_data.get('maquinario', {})
        frota_est = frota_data.get('frota_estimada_total', {})
        elements.append(Paragraph(f"<b>FROTA ESTIMADA:</b> {frota_est.get('tratores', 'N/D')} Tratores | {frota_est.get('colheitadeiras', 'N/D')} Colheitadeiras | Valor: {frota_data.get('valor_estimado_frota', 'N/D')}", body_style))
        
        elements.append(Spacer(1, 0.2*cm))
        
        # Conectividade
        conectividade = dossie_data.get('conectividade', {})
        elementos_municipios = conectividade.get('analise_por_municipio', [])
        if elementos_municipios:
            elements.append(Paragraph(f"<b>CONECTIVIDADE (4G/5G):</b> Analisados {len(elementos_municipios)} municípios", body_style))
            for mun in elementos_municipios[:3]:
                recomendacoes = mun.get('recomendacoes', [])
                rec_str = "; ".join(recomendacoes) if recomendacoes else "Boa cobertura"
                elements.append(Paragraph(f"• {mun.get('municipio')}, {mun.get('uf')}: {rec_str}", body_style))
        
        elements.append(Spacer(1, 0.4*cm))
        
        # ===== SEÇÃO 2: FINANCEIRO =====
        elements.append(PageBreak())
        elements.append(Paragraph("2️⃣ CAMADA FINANCEIRA & JURÍDICA (FOLLOW THE MONEY)", heading2_style))
        elements.append(Spacer(1, 0.3*cm))
        
        # CRA
        cra_data = dossie_data.get('cra_debentures', {})
        faturamento_real = cra_data.get('faturamento_real', 'N/D')
        ebitda = cra_data.get('ebitda_consolidado', 'N/D')
        alavancagem = cra_data.get('indice_dps', 'N/D')
        
        elementos.append(Paragraph(f"<b>DADOS AUDITADOS (CRA/Debêntures):</b>", body_style))
        elementos.append(Paragraph(f"Faturamento Real: {faturamento_real} | EBITDA: {ebitda} | D/EBITDA: {alavancagem}x", body_style))
        
        elements.append(Spacer(1, 0.2*cm))
        
        # Incentivos Fiscais
        incentivos = dossie_data.get('incentivos_fiscais', {})
        economia_anual = incentivos.get('economia_fiscal_anual_total', 'N/D')
        elements.append(Paragraph(f"<b>INCENTIVOS FISCAIS:</b> Economia anual: {economia_anual}", body_style))
        
        elements.append(Spacer(1, 0.2*cm))
        
        # Multas Ambientais
        multas = dossie_data.get('multas_ambientais', {})
        debitos_ambientais = multas.get('debitos_ambientais_total', 'N/D')
        score_risco_ambiental = multas.get('score_risco_ambiental', 'N/D')
        elements.append(Paragraph(f"<b>RISCOS AMBIENTAIS:</b> Débitos: {debitos_ambientais} | Score: {score_risco_ambiental}", body_style))
        
        elements.append(Spacer(1, 0.2*cm))
        
        # Processos Trabalhistas
        trt = dossie_data.get('processos_trabalhistas', {})
        total_procs = trt.get('total_processos_ativos', 0)
        horas_extras_pct = 0
        procs_por_tipo = trt.get('processos_por_tipo', {})
        if procs_por_tipo.get('horas_extras'):
            horas_extras_pct = (procs_por_tipo['horas_extras'] / total_procs * 100) if total_procs > 0 else 0
        elements.append(Paragraph(f"<b>PROCESSOS TRABALHISTAS:</b> {total_procs} processos ativos | {horas_extras_pct:.0f}% horas extras (ALERTA: falha em sistema de ponto)", body_style))
        
        elements.append(Spacer(1, 0.4*cm))
        
        # ===== SEÇÃO 3: SUPPLY CHAIN =====
        elements.append(PageBreak())
        elements.append(Paragraph("3️⃣ CADEIA DE SUPRIMENTOS", heading2_style))
        elements.append(Spacer(1, 0.3*cm))
        
        exportacao = dossie_data.get('exportacao', {})
        volume_exportado = exportacao.get('volume_total_exportado_2024', 'N/D')
        receita_export = exportacao.get('receita_exportacao_2024', 'N/D')
        elements.append(Paragraph(f"<b>EXPORTAÇÃO (2024):</b> {volume_exportado} | Receita: {receita_export}", body_style))
        
        elements.append(Spacer(1, 0.2*cm))
        
        bioinsumos = dossie_data.get('bioinsumos', {})
        total_biofabricas = bioinsumos.get('total_biofabricas', 0)
        capacidade_producao = bioinsumos.get('capacidade_producao_total', 'N/D')
        maturidade_bio = bioinsumos.get('maturidade_bioinsumos', 'N/D')
        elements.append(Paragraph(f"<b>BIOINSUMOS:</b> {total_biofabricas} biofábricas | Capacidade: {capacidade_producao} | Maturidade: {maturidade_bio}", body_style))
        
        elements.append(Spacer(1, 0.4*cm))
        
        # ===== SEÇÃO 4: TECNOLOGIA & PESSOAS =====
        elements.append(PageBreak())
        elements.append(Paragraph("4️⃣ TECNOLOGIA & PESSOAS (SOFT ASSETS)", heading2_style))
        elements.append(Spacer(1, 0.3*cm))
        
        tech_stack = dossie_data.get('tech_stack_identificado', {})
        stack_consolidado = tech_stack.get('stack_consolidado', {})
        elements.append(Paragraph(f"<b>TECH STACK:</b>", body_style))
        elements.append(Paragraph(f"ERP: {stack_consolidado.get('ERP_Principal', 'N/D')} | BD: {', '.join(stack_consolidado.get('Banco_Dados', []))} | BI: {', '.join(stack_consolidado.get('BI_Analytics', []))}", body_style))
        elements.append(Paragraph(f"Maturidade: {tech_stack.get('maturidade_ti', 'N/D')}", body_style))
        
        elements.append(Spacer(1, 0.2*cm))
        
        # E-mails
        emails_data = dossie_data.get('emails_decisores', {})
        total_emails = emails_data.get('total_emails_validados', 0)
        elements.append(Paragraph(f"<b>CONTATOS VALIDADOS:</b> {total_emails} e-mails de decisores identificados", body_style))
        
        emails_list = emails_data.get('emails_validos_identificados', [])
        if emails_list:
            for email_info in emails_list[:5]:
                nome = email_info.get('nome', '-')
                cargo = email_info.get('cargo', '-')
                email = email_info.get('email', '-')
                elements.append(Paragraph(f"• {nome} ({cargo}): {email}", body_style))
        
        elements.append(Spacer(1, 0.2*cm))
        
        # Funcionários
        funcionarios = dossie_data.get('estimativa_funcionarios', {})
        total_func = funcionarios.get('total_funcionarios_estimado', 'N/D')
        elements.append(Paragraph(f"<b>FORÇA DE TRABALHO:</b> ~{total_func} funcionários (estimado)", body_style))
        
        elements.append(Spacer(1, 0.4*cm))
        
        # ===== SEÇÃO 5: ANÁLISE ESTRATÉGICA =====
        elements.append(PageBreak())
        elements.append(Paragraph("5️⃣ ANÁLISE ESTRATÉGICA & PLANO DE ATAQUE", heading2_style))
        elements.append(Spacer(1, 0.3*cm))
        
        analise = dossie_data.get('analise_estrategica', {})
        
        secoes_analise = [
            ("QUEM É ESTA EMPRESA?", analise.get('quem_e_empresa', '')),
            ("DORES & COMPLEXIDADE", analise.get('complexidade_dores', '')),
            ("ARSENAL RECOMENDADO", analise.get('arsenal_recomendado', '')),
            ("PLANO DE ATAQUE", analise.get('plano_ataque', ''))
        ]
        
        for titulo_secao, conteudo_secao in secoes_analise:
            elements.append(Paragraph(f"<b>{titulo_secao}:</b>", heading2_style))
            # Limpa markdown e HTML
            conteudo_limpo = conteudo_secao.replace('**', '<b>').replace('**', '</b>').replace('##', '')
            conteudo_limpo = conteudo_limpo[:500] + "..." if len(conteudo_limpo) > 500 else conteudo_limpo  # Limita por espaço
            elements.append(Paragraph(conteudo_limpo, body_style))
            elements.append(Spacer(1, 0.3*cm))
        
        # ===== RODAPÉ =====
        elements.append(PageBreak())
        elements.append(Paragraph("=" * 80, body_style))
        elements.append(Paragraph("RELATÓRIO CONFIDENCIAL | USO RESTRITO SENIOR AGRO", body_style))
        elements.append(Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}", body_style))
        
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def generate_docx(dossie_data):
        """Gera DOCX editável com todas as camadas."""
        doc = Document()
        
        # Estilos
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(10)
        
        # Cabeçalho
        titulo = doc.add_heading(f"🔴 DOSSIÊ FOX-3: {dossie_data.get('empresa_alvo', 'ALVO').upper()}", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitulo = doc.add_paragraph()
        subtitulo.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')} | Score SAS: {dossie_data.get('sas_score', 'N/A')} ({dossie_data.get('sas_tier', 'N/A')})").bold = True
        
        # Tabela Resumo
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'INDICADOR'
        hdr_cells[1].text = 'VALOR'
        
        resumo = {
            'Área Total': f"{dossie_data.get('dados_operacionais', {}).get('area_total', 'N/D')} ha",
            'Faturamento': dossie_data.get('dados_financeiros', {}).get('faturamento_estimado', 'N/D'),
            'EBITDA': dossie_data.get('dados_financeiros', {}).get('ebitda_ajustado', 'N/D'),
            'ERP': dossie_data.get('tech_stack', {}).get('erp_principal', 'N/D'),
            'Funcionários': dossie_data.get('dados_operacionais', {}).get('funcionarios_estimados', 'N/D')
        }
        
        for chave, valor in resumo.items():
            linha = table.add_row().cells
            linha[0].text = chave
            linha[1].text = str(valor)
        
        doc.add_paragraph("")
        
        # Seções principais
        doc.add_heading("1. INFRAESTRUTURA (Hard Assets)", level=1)
        doc.add_paragraph(f"Área SIGEF/CAR: {dossie_data.get('sigef_car', {}).get('area_total_hectares', 'N/D')} ha")
        doc.add_paragraph(f"Frota Estimada: {dossie_data.get('maquinario', {}).get('frota_estimada_total', {}).get('tratores', 'N/D')} tratores, {dossie_data.get('maquinario', {}).get('frota_estimada_total', {}).get('colheitadeiras', 'N/D')} colheitadeiras")
        
        doc.add_heading("2. FINANCEIRO (Follow the Money)", level=1)
        cra = dossie_data.get('cra_debentures', {})
        doc.add_paragraph(f"Faturamento Real (CRA): {cra.get('faturamento_real', 'N/D')}")
        doc.add_paragraph(f"EBITDA Consolidado: {cra.get('ebitda_consolidado', 'N/D')}")
        doc.add_paragraph(f"Alavancagem (D/EBITDA): {cra.get('indice_dps', 'N/D')}x")
        
        doc.add_heading("3. CADEIA DE SUPRIMENTOS", level=1)
        doc.add_paragraph(f"Exportação 2024: {dossie_data.get('exportacao', {}).get('volume_total_exportado_2024', 'N/D')}")
        doc.add_paragraph(f"Biofábricas: {dossie_data.get('bioinsumos', {}).get('total_biofabricas', 0)} unidades")
        
        doc.add_heading("4. TECNOLOGIA & PESSOAS", level=1)
        tech = dossie_data.get('tech_stack_identificado', {})
        doc.add_paragraph(f"ERP: {tech.get('stack_consolidado', {}).get('ERP_Principal', 'N/D')}")
        doc.add_paragraph(f"Contatos Identificados: {dossie_data.get('emails_decisores', {}).get('total_emails_validados', 0)}")
        
        doc.add_heading("5. ANÁLISE ESTRATÉGICA", level=1)
        analise = dossie_data.get('analise_estrategica', {})
        doc.add_paragraph(analise.get('quem_e_empresa', 'Sem análise'))
        
        doc.add_paragraph("")
        doc.add_paragraph("=" * 80)
        doc.add_paragraph(f"Confidencial | Gerado {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
