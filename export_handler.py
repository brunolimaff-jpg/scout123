"""
utils/export_handler.py — Sistema de Exportação Defensivo
CORRIGIDO: Importações opcionais com fallback gracioso
"""
import io
from datetime import datetime
import json
import logging

logger = logging.getLogger(__name__)

# ========== IMPORTS OPCIONAIS (ReportLab) ==========
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.units import cm
    HAS_REPORTLAB = True
    logger.info("[ExportHandler] ReportLab disponível")
except ImportError as e:
    HAS_REPORTLAB = False
    logger.warning(f"[ExportHandler] ReportLab NÃO instalado: {e}")
    # Placeholders para evitar erros
    A4 = None
    colors = None

# ========== IMPORTS OPCIONAIS (python-docx) ==========
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
    logger.info("[ExportHandler] python-docx disponível")
except ImportError as e:
    HAS_DOCX = False
    logger.warning(f"[ExportHandler] python-docx NÃO instalado: {e}")
    Document = None


class ExportHandler:
    """
    Handler para exportação de dossiês em múltiplos formatos.
    Funciona mesmo sem dependências instaladas (com mensagens de erro).
    """
    
    @staticmethod
    def generate_pdf(dossie_data):
        """
        Gera PDF profissional do dossiê RADAR FOX-3.
        
        Args:
            dossie_data: Dicionário com dados completos do dossiê
            
        Returns:
            BytesIO buffer com PDF ou levanta exceção se ReportLab não instalado
        """
        if not HAS_REPORTLAB:
            raise ImportError(
                "📦 ReportLab não está instalado!\n\n"
                "Para gerar PDFs, adicione ao requirements.txt:\n"
                "reportlab>=4.0.0\n\n"
                "OU use exportação em DOCX/JSON."
            )
        
        logger.info("[ExportHandler] Gerando PDF...")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=A4,
            rightMargin=1.5*cm, 
            leftMargin=1.5*cm,
            topMargin=1.5*cm, 
            bottomMargin=1.5*cm
        )
        
        elements = []
        styles = getSampleStyleSheet()
        
        # ===== ESTILOS =====
        title_style = ParagraphStyle(
            'TitleCustom', 
            parent=styles['Heading1'], 
            fontSize=22, 
            spaceAfter=12, 
            textColor=colors.HexColor("#0f172a"), 
            alignment=1, 
            fontName='Helvetica-Bold'
        )
        
        heading2_style = ParagraphStyle(
            'Heading2Custom', 
            parent=styles['Heading2'], 
            fontSize=14, 
            spaceAfter=8,
            textColor=colors.HexColor("#1e40af"), 
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'BodyCustom', 
            parent=styles['Normal'], 
            fontSize=9, 
            spaceAfter=6, 
            leading=12
        )
        
        # ===== CABEÇALHO =====
        elements.append(Paragraph("🔴 RADAR FOX-3 | INTELLIGENCE REPORT", title_style))
        elements.append(Spacer(1, 0.3*cm))
        
        empresa_alvo = dossie_data.get('empresa_alvo', 'ALVO DESCONHECIDO')
        sas_score = dossie_data.get('sas_score', 'N/A')
        sas_tier = dossie_data.get('sas_tier', 'N/A')
        
        elements.append(Paragraph(
            f"<b>ALVO:</b> {empresa_alvo.upper()} | "
            f"<b>SCORE SAS:</b> {sas_score} ({sas_tier}) | "
            f"<b>DATA:</b> {datetime.now().strftime('%d/%m/%Y %H:%M')}", 
            body_style
        ))
        elements.append(Spacer(1, 0.5*cm))
        
        # ===== SCORECARD RESUMIDO =====
        elementos_resumo = [
            ['INDICADOR', 'DADOS RECUPERADOS'],
            ['Área Total', f"{dossie_data.get('dados_operacionais', {}).get('area_total', 'N/D')} ha"],
            ['Faturamento', dossie_data.get('dados_financeiros', {}).get('faturamento_estimado', 'N/D')],
            ['ERP Principal', dossie_data.get('tech_stack', {}).get('erp_principal', 'N/D')],
        ]
        
        table_resumo = Table(elementos_resumo, colWidths=[4*cm, 8*cm])
        table_resumo.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
        ]))
        elements.append(table_resumo)
        elements.append(Spacer(1, 0.5*cm))
        
        # ===== ANÁLISE ESTRATÉGICA =====
        elements.append(PageBreak())
        elements.append(Paragraph("🎯 ANÁLISE ESTRATÉGICA", heading2_style))
        
        analise = dossie_data.get('analise_estrategica', {})
        
        for titulo_secao, conteudo_secao in [
            ("QUEM É ESTA EMPRESA?", analise.get('quem_e_empresa', 'Análise indisponível')),
            ("DORES & COMPLEXIDADE", analise.get('complexidade_dores', 'Análise indisponível')),
            ("ARSENAL RECOMENDADO", analise.get('arsenal_recomendado', 'Análise indisponível')),
            ("PLANO DE ATAQUE", analise.get('plano_ataque', 'Análise indisponível'))
        ]:
            elements.append(Paragraph(f"<b>{titulo_secao}:</b>", heading2_style))
            # Limpa markdown
            conteudo_limpo = conteudo_secao.replace('**', '').replace('##', '')[:500]
            elements.append(Paragraph(conteudo_limpo, body_style))
            elements.append(Spacer(1, 0.3*cm))
        
        # ===== RODAPÉ =====
        elements.append(PageBreak())
        elements.append(Paragraph("=" * 80, body_style))
        elements.append(Paragraph("RELATÓRIO CONFIDENCIAL | USO RESTRITO SENIOR AGRO", body_style))
        elements.append(Paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}", body_style))
        
        doc.build(elements)
        buffer.seek(0)
        
        logger.info("[ExportHandler] PDF gerado com sucesso")
        return buffer
    
    @staticmethod
    def generate_docx(dossie_data):
        """
        Gera DOCX editável do dossiê.
        
        Args:
            dossie_data: Dicionário com dados completos do dossiê
            
        Returns:
            BytesIO buffer com DOCX ou levanta exceção se python-docx não instalado
        """
        if not HAS_DOCX:
            raise ImportError(
                "📦 python-docx não está instalado!\n\n"
                "Para gerar DOCX, adicione ao requirements.txt:\n"
                "python-docx>=1.1.0\n\n"
                "OU use exportação em JSON."
            )
        
        logger.info("[ExportHandler] Gerando DOCX...")
        
        doc = Document()
        
        # Estilos
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(10)
        
        # Cabeçalho
        titulo = doc.add_heading(f"🔴 DOSSIÊ FOX-3: {dossie_data.get('empresa_alvo', 'ALVO').upper()}", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitulo = doc.add_paragraph()
        subtitulo.add_run(
            f"Data: {datetime.now().strftime('%d/%m/%Y')} | "
            f"Score SAS: {dossie_data.get('sas_score', 'N/A')} ({dossie_data.get('sas_tier', 'N/A')})"
        ).bold = True
        
        # Tabela Resumo
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'INDICADOR'
        hdr_cells[1].text = 'VALOR'
        
        resumo = {
            'Área Total': f"{dossie_data.get('dados_operacionais', {}).get('area_total', 'N/D')} ha",
            'Faturamento': dossie_data.get('dados_financeiros', {}).get('faturamento_estimado', 'N/D'),
            'ERP': dossie_data.get('tech_stack', {}).get('erp_principal', 'N/D'),
        }
        
        for chave, valor in resumo.items():
            linha = table.add_row().cells
            linha[0].text = chave
            linha[1].text = str(valor)
        
        doc.add_paragraph("")
        
        # Análise Estratégica
        doc.add_heading("🎯 ANÁLISE ESTRATÉGICA", level=1)
        analise = dossie_data.get('analise_estrategica', {})
        doc.add_paragraph(analise.get('quem_e_empresa', 'Análise indisponível'))
        
        # Rodapé
        doc.add_paragraph("")
        doc.add_paragraph("=" * 80)
        doc.add_paragraph(f"Confidencial | Gerado {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("[ExportHandler] DOCX gerado com sucesso")
        return buffer
    
    @staticmethod
    def generate_json(dossie_data):
        """
        Gera JSON bruto do dossiê (sempre disponível).
        
        Args:
            dossie_data: Dicionário com dados completos do dossiê
            
        Returns:
            BytesIO buffer com JSON
        """
        logger.info("[ExportHandler] Gerando JSON...")
        
        json_str = json.dumps(dossie_data, indent=2, ensure_ascii=False, default=str)
        buffer = io.BytesIO(json_str.encode('utf-8'))
        buffer.seek(0)
        
        logger.info("[ExportHandler] JSON gerado com sucesso")
        return buffer
